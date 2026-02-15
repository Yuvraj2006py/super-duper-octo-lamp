import textwrap
from pathlib import Path
from typing import Any

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas


def _replace_in_paragraph(paragraph, context: dict[str, Any]) -> None:
    for key, value in context.items():
        placeholder = "{{" + key + "}}"
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, str(value))


def render_docx_template(template_path: Path, output_path: Path, context: dict[str, Any]) -> None:
    document = Document(str(template_path))

    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, context)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, context)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def render_text_pdf(output_path: Path, *, title: str, body: str) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    pdf = canvas.Canvas(str(output_path), pagesize=LETTER)
    width, height = LETTER
    margin = 54
    line_height = 14
    max_chars = 96
    y = height - margin

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(margin, y, title)
    y -= line_height * 2

    pdf.setFont("Helvetica", 11)
    for paragraph in (body or "").splitlines():
        lines = textwrap.wrap(paragraph, width=max_chars) or [""]
        for line in lines:
            if y < margin:
                pdf.showPage()
                pdf.setFont("Helvetica", 11)
                y = height - margin
            pdf.drawString(margin, y, line)
            y -= line_height
        y -= 4

    pdf.save()
